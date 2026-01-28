#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
测试脚本：用于测试 testfile 目录中的 docx 文件
该脚本将上传 docx 文件到 OntologySystem API 并测试处理流程
"""

import os
import requests
import json
from pathlib import Path
from typing import List, Dict, Optional

# 添加加载 .env 文件的功能
from dotenv import load_dotenv
load_dotenv()  # 加载 .env 文件到环境变量


def read_docx_as_text(filepath: str) -> str:
    """
    读取 docx 文件内容并转换为文本
    """
    try:
        from docx import Document
        print(f"正在使用 python-docx 读取文件: {filepath}")
        doc = Document(filepath)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 只添加非空段落
                text_parts.append(paragraph.text)
        
        full_text = "\n".join(text_parts)
        print(f"从文档中读取到 {len(text_parts)} 个段落，总字符数: {len(full_text)}")
        
        # 同时尝试读取表格中的内容
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_text.append(cell.text)
        
        if table_text:
            print(f"从表格中读取到 {len(table_text)} 项内容")
            full_text += "\n" + "\n".join(table_text)
        
        return full_text
    except ImportError:
        print("错误: 未安装 python-docx，请运行: pip install python-docx")
        return ""
    except Exception as e:
        print(f"读取 docx 文件时发生错误: {e}")
        return ""


def load_config_from_env():
    """
    从环境变量加载大模型配置
    """
    print("正在检查环境变量...")
    print(f"VLLM_API_KEY 存在: {os.getenv('VLLM_API_KEY') is not None}")
    print(f"VLLM_API_KEY 值: '{os.getenv('VLLM_API_KEY')}'")
    print(f"VLLM_BASE_URL 存在: {os.getenv('VLLM_BASE_URL') is not None}")
    print(f"VLLM_BASE_URL 值: '{os.getenv('VLLM_BASE_URL')}'")
    print(f"OPENROUTER_API_KEY 存在: {os.getenv('OPENROUTER_API_KEY') is not None}")
    print(f"OPENROUTER_API_KEY 值: '{os.getenv('OPENROUTER_API_KEY')}'")
    print(f"OPENROUTER_BASE_URL 存在: {os.getenv('OPENROUTER_BASE_URL') is not None}")
    print(f"OPENROUTER_BASE_URL 值: '{os.getenv('OPENROUTER_BASE_URL')}'")

    # 优先使用 vLLM 配置（只要 BASE_URL 存在且非空）
    vllm_api_key = os.getenv("VLLM_API_KEY")
    vllm_base_url = os.getenv("VLLM_BASE_URL")
    vllm_model = os.getenv("VLLM_MODEL")
    if vllm_base_url and vllm_base_url.strip():
        print("使用 vLLM 配置")
        return {
            "api_key": vllm_api_key or "",  # 即使为空也使用
            "base_url": vllm_base_url,
            "model": vllm_model or "qwen2.5-7B"
        }
    
    # 其次使用 OpenRouter 配置
    openrouter_api_key = os.getenv("OPENROUTER_API_KEY")
    openrouter_base_url = os.getenv("OPENROUTER_BASE_URL")
    openrouter_model = os.getenv("OPENROUTER_MODEL")
    if openrouter_api_key and openrouter_api_key.strip() and openrouter_base_url and openrouter_base_url.strip():
        print("使用 OpenRouter 配置")
        return {
            "api_key": openrouter_api_key,
            "base_url": openrouter_base_url,
            "model": openrouter_model or "z-ai/glm-4.5-air:free"
        }
    
    # 如果以上都没有，抛出错误
    raise ValueError("未找到有效的 LLM 配置，请检查 .env 文件")


def upload_and_test_file(filepath: str, api_base_url: str = "http://localhost:3001"):
    """
    上传文件并测试本体生成
    """
    print(f"开始测试文件: {filepath}")
    
    # 读取 docx 文件内容
    try:
        file_content = read_docx_as_text(filepath)
        
        if len(file_content.strip()) == 0:
            print("警告: 文件内容为空或无法读取")
            # 尝试备用方法：直接使用文件路径上传（如果API支持）
            print("尝试使用文件上传API...")
            upload_file_via_api(filepath, api_base_url)
            return
        else:
            print(f"成功读取文件，内容长度: {len(file_content)} 字符")
            print(f"内容预览 (前 500 字符): {file_content[:500]}...")
    except Exception as e:
        print(f"读取文件失败: {e}")
        return

    # 加载 LLM 配置
    llm_config = load_config_from_env()
    print(f"使用 LLM 配置: {llm_config['model']} 从 {llm_config['base_url']}")

    # 调用本体生成 API
    ontology_api_url = f"{api_base_url}/api/v1/ontology/generate"
    
    payload = {
        "text_content": file_content[:10000],  # 限制内容长度以避免超时
        "scenario": "",
        "rules": [
            {
                "cls_name": "技术与知识领域",
                "attrs": "描述, 成熟度",
                "rels": "支撑(创新载体), 应用于(业务价值)"
            },
            {
                "cls_name": "创新阶段",
                "attrs": "状态, 时间节点",
                "rels": "演进至(下一阶段), 包含(创新载体)"
            },
            {
                "cls_name": "创新载体与平台",
                "attrs": "版本, 功能列表",
                "rels": "承载(试点功能), 由...构建(组织)"
            },
            {
                "cls_name": "创新方法与机制",
                "attrs": "方法论描述",
                "rels": "应用于(项目), 产出(成果)"
            },
            {
                "cls_name": "组织主体",
                "attrs": "角色, 职责",
                "rels": "合作(其他组织), 主导(创新载体)"
            },
            {
                "cls_name": "业务价值目标",
                "attrs": "预期成效, 指标",
                "rels": "关联(技术), 属于(业务条线)"
            }
        ],
        "api_key": llm_config["api_key"],
        "base_url": llm_config["base_url"],
        "model": llm_config["model"]
    }
    
    headers = {
        "Content-Type": "application/json"
    }
    
    print("正在发送请求到本体生成API...")
    print(f"请求大小: {len(json.dumps(payload))} 字节")
    try:
        response = requests.post(ontology_api_url, headers=headers, json=payload, timeout=600)  # 增加超时到600秒
        
        if response.status_code == 200:
            result = response.json()
            print("✅ 本体生成成功!")
            print(f"文件名: {result.get('filename', 'N/A')}")
            print(f"处理日志: {result.get('log', 'N/A')}")
            
            # 保存生成的本体文件
            ttl_content = result.get('ttl_content', '')
            if ttl_content:
                output_filename = f"generated_ontology_{Path(filepath).stem}.ttl"
                with open(output_filename, 'w', encoding='utf-8') as f:
                    f.write(ttl_content)
                print(f"TTL 本体文件已保存到: {output_filename}")
        else:
            print(f"❌ 本体生成失败，状态码: {response.status_code}")
            print(f"错误详情: {response.text}")
            
    except requests.exceptions.Timeout:
        print("❌ 请求超时")
    except requests.exceptions.RequestException as e:
        print(f"❌ 请求失败: {e}")


def upload_file_via_api(filepath: str, api_base_url: str = "http://localhost:3001"):
    """
    如果直接读取失败，尝试使用文件上传API（如果存在）
    """
    # 尝试查找文件上传的API端点
    files_api_urls = [
        f"{api_base_url}/api/v1/files/upload",
        f"{api_base_url}/upload",
        f"{api_base_url}/api/v1/files"
    ]
    
    for upload_url in files_api_urls:
        try:
            with open(filepath, 'rb') as f:
                files = {'file': (Path(filepath).name, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                response = requests.post(upload_url, files=files, timeout=300)
                
                if response.status_code in [200, 201]:
                    print(f"✅ 文件上传成功到 {upload_url}")
                    print(f"响应: {response.text}")
                    return
        except Exception as e:
            print(f"尝试上传到 {upload_url} 失败: {e}")
            continue
    
    print("❌ 所有上传端点都不可用或失败")


def test_all_files_in_directory(directory: str, api_base_url: str = "http://localhost:3001"):
    """
    测试目录中的所有文件
    """
    directory_path = Path(directory)
    
    if not directory_path.exists():
        print(f"目录不存在: {directory}")
        return
    
    # 支持的文件扩展名
    supported_extensions = {'.docx', '.txt', '.pdf', '.md'}
    
    files_to_process = []
    for file_path in directory_path.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            files_to_process.append(str(file_path))
    
    if not files_to_process:
        print(f"在目录 {directory} 中未找到支持的文件类型 {supported_extensions}")
        return
    
    print(f"找到 {len(files_to_process)} 个待处理文件")
    
    for filepath in files_to_process:
        print(f"\n--- 处理文件: {Path(filepath).name} ---")
        upload_and_test_file(filepath, api_base_url)
        print("-" * 50)


def main():
    """
    主函数
    """
    test_directory = "/home/lenovo/Documents/PythonProject/OntologySystem/testfile"
    api_base_url = "http://localhost:3001"  # 默认API地址
    
    print("开始测试 testfile 目录中的文件...")
    print(f"测试目录: {test_directory}")
    print(f"API 地址: {api_base_url}")
    
    test_all_files_in_directory(test_directory, api_base_url)
    

if __name__ == "__main__":
    main()